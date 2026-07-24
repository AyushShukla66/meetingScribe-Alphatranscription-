from pathlib import Path
from docx import Document


def generate_report(meeting):

    Path("reports").mkdir(exist_ok=True)

    doc = Document()

    doc.add_heading("Finodaya Capital", level=1)
    doc.add_heading("Executive Meeting Brief", level=2)

    doc.add_paragraph(f"Meeting : {meeting.name}")
    doc.add_paragraph(f"Start   : {meeting.start_time}")
    doc.add_paragraph(f"End     : {meeting.end_time}")

    doc.save(f"reports/{meeting.name}.docx")