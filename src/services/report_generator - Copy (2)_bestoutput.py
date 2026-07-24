from pathlib import Path
from docx import Document

from services.transcription import transcribe_audio
from services.ai_engine import generate_executive_brief


def generate_report(meeting):

    Path("reports").mkdir(exist_ok=True)

    # 1. Convert audio to text
    transcript = transcribe_audio(
        "temp/meeting.wav"
    )

    print("======== TRANSCRIPT ========")
    print(transcript)
    print("============================")


    # 2. Generate AI executive brief
    brief = generate_executive_brief(
        transcript
    )


    # 3. Create Word report
    doc = Document()

    doc.add_heading("Finodaya Capital", level=1)
    doc.add_heading("Executive Meeting Brief", level=2)

    doc.add_paragraph(
        f"Meeting : {meeting.name}"
    )

    doc.add_paragraph(
        f"Start : {meeting.start_time}"
    )

    doc.add_paragraph(
        f"End : {meeting.end_time}"
    )


    doc.add_heading("AI Generated Brief", level=2)

    doc.add_paragraph(
        str(brief)
    )


    doc.add_heading("Transcript", level=2)

    doc.add_paragraph(
        transcript
    )


    filename = f"reports/{meeting.name}.docx"

    doc.save(filename)

    print("REPORT SAVED:", Path(filename).absolute())

    return filename