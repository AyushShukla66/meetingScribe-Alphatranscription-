import re

from pathlib import Path

from docx import Document
from openai import OpenAI

from config.settings import OPENAI_API_KEY
from services.transcription import transcribe_audio
from services.ai_engine import generate_executive_brief


client = OpenAI(
    api_key=OPENAI_API_KEY,
    timeout=300.0
)



def create_meeting_report(meeting, meeting_report, safe_name):
    """
    Create AI generated meeting report DOCX.
    """

    doc = Document()


    doc.add_heading(
        "Finodaya Capital",
        level=1
    )


    doc.add_heading(
        "Meeting Report",
        level=2
    )


    doc.add_paragraph(
        f"Meeting Name: {meeting.name}"
    )


    doc.add_paragraph(
        f"Start Time: {meeting.start_time}"
    )


    doc.add_paragraph(
        f"End Time: {meeting.end_time}"
    )


    doc.add_heading(
        "Summary",
        level=2
    )


    doc.add_paragraph(
        meeting_report
    )


    filename = (
        f"reports/{safe_name}_Meeting_Report.docx"
    )


    doc.save(filename)


    print(
        "Meeting Report Created:",
        Path(filename).absolute()
    )


    return filename





def create_transcript_report(meeting, transcript, safe_name):
    """
    Create original voice transcript DOCX.
    """


    doc = Document()


    doc.add_heading(
        "Finodaya Capital",
        level=1
    )


    doc.add_heading(
        "Original Voice Transcript",
        level=2
    )


    doc.add_paragraph(
        f"Meeting Name: {meeting.name}"
    )


    doc.add_paragraph(
        f"Start Time: {meeting.start_time}"
    )


    doc.add_paragraph(
        f"End Time: {meeting.end_time}"
    )


    doc.add_heading(
        "Transcript",
        level=2
    )


    doc.add_paragraph(
        transcript
    )


    filename = (
        f"reports/{safe_name}_Original_Transcript.docx"
    )


    doc.save(filename)


    print(
        "Original Transcript Created:",
        Path(filename).absolute()
    )


    return filename





def create_executive_brief(meeting, executive_brief, safe_name):
    """
    Create Executive Brief DOCX.
    """


    doc = Document()


    doc.add_heading(
        "Finodaya Capital",
        level=1
    )


    doc.add_heading(
        "Executive Brief",
        level=2
    )


    doc.add_paragraph(
        f"Meeting Name: {meeting.name}"
    )


    doc.add_paragraph(
        f"Start Time: {meeting.start_time}"
    )


    doc.add_paragraph(
        f"End Time: {meeting.end_time}"
    )


    doc.add_heading(
        "Executive Summary",
        level=2
    )


    doc.add_paragraph(
        executive_brief
    )


    filename = (
        f"reports/{safe_name}_Executive_Brief.docx"
    )


    doc.save(filename)


    print(
        "Executive Brief Created:",
        Path(filename).absolute()
    )


    return filename





def validate_transcript(transcript):
    """
    Check if transcript is valid.
    """


    if transcript is None:

        return False



    if not isinstance(transcript, str):

        return False



    cleaned = transcript.strip()



    if len(cleaned) < 20:

        return False



    return True





def generate_report(meeting):
    """
    Generate meeting report,
    transcript report,
    and executive brief.
    """


    Path("reports").mkdir(
        exist_ok=True
    )



    print(
        "Generating transcript..."
    )



    if not hasattr(meeting, "audio_file"):

        raise Exception(
            "Meeting audio file not found."
        )



    if not Path(meeting.audio_file).exists():

        raise FileNotFoundError(
            meeting.audio_file
        )



    transcript = transcribe_audio(
        meeting.audio_file
    )



    if not validate_transcript(transcript):

        raise Exception(
            "Transcript generation failed or transcript is too short."
        )



    # Save transcript in meeting object

    meeting.transcript = transcript



    print(
        "\n======== ORIGINAL TRANSCRIPT ========\n"
    )


    print(
        transcript
    )



    with open(
        "temp/raw_transcript.txt",
        "w",
        encoding="utf-8"
    ) as file:

        file.write(
            transcript
        )



    print(
        "Generating meeting summary..."
    )



    report_response = client.chat.completions.create(

        model="gpt-4o-mini",

        messages=[

            {
                "role": "system",
                "content": """

You are a professional corporate meeting report generator.

Create a structured meeting report from the transcript.

Generate only:

1. Meeting Topic
2. Key Discussion Points
3. Decisions Made
4. Action Items


Rules:

- Write the report in professional English.
- Do not include the transcript.
- Do not modify the transcript.
- Do not add information not discussed.
- Keep technical terms unchanged.
- Keep API names unchanged.
- Keep software names unchanged.
- Keep company names unchanged.
- Keep product names unchanged.

"""
            },


            {
                "role": "user",
                "content": transcript
            }

        ]
    )



    meeting_report = (
        report_response
        .choices[0]
        .message
        .content
    )



    safe_name = re.sub(
        r'[\\/:*?"<>|]',
        "-",
        meeting.name
    )



    meeting_report_file = create_meeting_report(
        meeting,
        meeting_report,
        safe_name
    )



    transcript_file = create_transcript_report(
        meeting,
        transcript,
        safe_name
    )



    print(
        "Generating executive brief..."
    )



    executive_brief = generate_executive_brief(
        transcript
    )



    executive_brief_file = create_executive_brief(
        meeting,
        executive_brief,
        safe_name
    )



    print(
        "\nAll documents generated successfully."
    )


    print(
        "Meeting Report:",
        meeting_report_file
    )


    print(
        "Transcript:",
        transcript_file
    )


    print(
        "Executive Brief:",
        executive_brief_file
    )



    return meeting_report_file